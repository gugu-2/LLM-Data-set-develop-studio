"""
Hypasia AI — DOCX parser using python-docx.
Extracts paragraphs and headings, groups into sections.
"""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from hypasia.schema import HypasiaRow


def parse_docx(path: Path) -> list[HypasiaRow]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed: pip install python-docx")

    doc = Document(path)
    source = str(path.resolve())
    rows: list[HypasiaRow] = []

    # Group paragraphs into sections by headings
    sections: list[tuple[str, str]] = []  # (heading, body)
    current_heading = path.stem
    current_body: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            if current_body:
                sections.append((current_heading, "\n".join(current_body)))
            current_heading = text
            current_body = []
        else:
            current_body.append(text)

    if current_body:
        sections.append((current_heading, "\n".join(current_body)))

    if not sections:
        # Fallback: dump everything
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if all_text:
            rows.append(HypasiaRow(
                instruction=f"What does the document '{path.name}' cover?",
                response=all_text,
                source=source,
                source_type="file",
                title=path.stem,
                raw_text=all_text,
                date_extracted=datetime.now(timezone.utc).isoformat(),
            ))
        return rows

    for heading, body in sections:
        if len(body.split()) < 10:
            continue
        rows.append(HypasiaRow(
            instruction=f"Explain the section '{heading}' from {path.name}.",
            response=body,
            source=source,
            source_type="file",
            title=heading,
            raw_text=body,
            date_extracted=datetime.now(timezone.utc).isoformat(),
        ))

    return rows
